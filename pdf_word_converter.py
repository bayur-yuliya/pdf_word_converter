import logging
import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml

# Setting up logging
logging.basicConfig(filename='pdf_to_docx.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def convert_pdf_to_docx(pdf_file, docx_file):
    try:
        pdf = pdfplumber.open(pdf_file)
        doc = Document()

        for page in pdf.pages:
            # Extracting text and hyperlinks from each page
            text = page.extract_text()
            hyperlinks = page.hyperlinks

            # Add text with hyperlinks to a Word document
            for line in text.split('\n'):
                paragraph = doc.add_paragraph(line)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(6)
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                # Adding active links
                for hyperlink in hyperlinks:
                    if hyperlink['bottom'] > page.height * 0.5 and hyperlink['top'] < page.height * 0.5:
                        run = paragraph.add_run()
                        run.add_text(hyperlink['uri'])
                        link = parse_xml(r'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:rPr><w:color w:val="0000FF"/></w:rPr><w:t> </w:t></w:r><w:t>{0}</w:t></w:hyperlink>'.format(hyperlink['uri']))
                        run._r.append(link)

        doc.save(docx_file)
        logging.info(f'PDF-файл "{pdf_file}" успешно преобразован в документ Word "{docx_file}".')
    except Exception as e:
        logging.error(f'Ошибка при преобразовании PDF в DOCX: {e}')


if __name__ == "__main__":
    pdf_file = input("Введите имя PDF-файла для преобразования: ") + '.pdf'
    docx_file = input("Введите имя DOCX-файла для сохранения результата: ") + '.docx'
    convert_pdf_to_docx(pdf_file, docx_file)
